"""
resume_build_script.py — python-docx equivalent of resume_build_script.js

WHY THIS VERSION EXISTS:
Some sandboxes (e.g. ChatGPT's Code Interpreter) have no network access, so
`npm install docx` fails and the Node.js script cannot run. python-docx is
usually already installed in those environments. This script produces the
same 2-page layout using the exact same spacing/color/size constants as the
Node.js version — do not change them unless the user explicitly asks for a
different visual style.

============================================================================
INSTRUCTIONS FOR ANY LLM ASKED TO ADAPT THIS SCRIPT FOR NEW RESUME CONTENT:

1. RUN THIS SCRIPT AS-IS. Do not rewrite the styling/formatting logic. Only
   edit the CONTENT section near the bottom (inside build_resume()) to swap
   in new text — header info, bullets, roles, etc.

2. BOLD METRICS RULE: every quantified achievement (dollar amounts,
   percentages, multipliers like "15x", counts like "14+ markets") must be
   its own bold run within the bullet, exactly like the existing bullets
   below via add_bullet(paragraph_segments). Do this for EVERY bullet in
   EVERY role — not just the first one. If unsure whether a number is a
   "metric," bold it.

3. DO NOT insert empty paragraphs as spacers anywhere. All spacing must come
   from paragraph_format.space_before / space_after — never from blank
   paragraphs.

4. DO NOT let python-docx's default paragraph spacing/line spacing apply.
   Always explicitly set space_before, space_after, and line_spacing on every
   paragraph via the helper functions below. Line spacing must be exactly
   1.0 (single) — never leave it unset (Word's Normal style default is often
   1.08, which will silently blow past 2 pages).

5. AFTER GENERATING: convert to PDF (e.g. via LibreOffice headless:
   `soffice --headless --convert-to pdf resume.docx`) and count pages. If
   it's not exactly 2, reduce the SECTION_SPACE_* / BULLET_SPACE_AFTER
   constants uniformly by ~15-20% and regenerate — never fix overflow by
   cutting content or by changing just one section's spacing.
============================================================================
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Design tokens (KEEP IDENTICAL TO resume_build_script.js) ----------
NAVY = RGBColor(0x1F, 0x38, 0x64)
DARKGRAY = RGBColor(0x33, 0x33, 0x33)
MIDGRAY = RGBColor(0x55, 0x55, 0x55)
FONT = "Calibri"

BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(9)
NAME_SIZE = Pt(20)
TITLE_SIZE = Pt(10.5)
H2_SIZE = Pt(10.5)

SECTION_SPACE_BEFORE = Pt(5.2)   # identical before every section header, no exceptions
SECTION_SPACE_AFTER = Pt(2.4)    # identical after every section header, no exceptions
BULLET_SPACE_AFTER = Pt(1.2)     # identical after every bullet, no exceptions
ROLE_SPACE_BEFORE = Pt(5.2)
ROLE_SPACE_AFTER = Pt(1.2)
LINE_SPACING = 1.0               # single — never 1.08/1.15


# ---------- Low-level helpers ----------
def set_run(run, size=BODY_SIZE, color=DARKGRAY, bold=False, italic=False, font=FONT):
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    return run


def set_paragraph_spacing(p, before=Pt(0), after=Pt(0), line=LINE_SPACING):
    pf = p.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line
    return p


def add_bottom_border(paragraph, color="1F3864", sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=SECTION_SPACE_BEFORE, after=SECTION_SPACE_AFTER)
    r = p.add_run(text.upper())
    set_run(r, size=H2_SIZE, color=NAVY, bold=True)
    add_bottom_border(p)
    return p


def add_bullet(doc, segments):
    """segments: list of (text, bold) tuples"""
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=Pt(0), after=BULLET_SPACE_AFTER)
    pf = p.paragraph_format
    pf.left_indent = Pt(13.7)
    pf.first_line_indent = Pt(-13.7)
    for text, bold in segments:
        r = p.add_run(text)
        set_run(r, size=BODY_SIZE, color=DARKGRAY, bold=bold)
    return p


def add_role_header(doc, title, company, dates):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=ROLE_SPACE_BEFORE, after=ROLE_SPACE_AFTER)
    # right tab stop at right margin (6.5in usable width w/ 0.5in margins on Letter = 7.5in text width)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    r1 = p.add_run(company)
    set_run(r1, size=Pt(BODY_SIZE.pt + 1), color=NAVY, bold=True)
    r2 = p.add_run(f"  |  {title}")
    set_run(r2, size=BODY_SIZE, color=MIDGRAY, italic=True)
    r3 = p.add_run("\t" + dates)
    set_run(r3, size=SMALL_SIZE, color=MIDGRAY, italic=True)
    return p


def add_sub_label(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(3.5))
    r1 = p.add_run(f"{label}: ")
    set_run(r1, size=BODY_SIZE, color=NAVY, bold=True)
    r2 = p.add_run(value)
    set_run(r2, size=BODY_SIZE, color=DARKGRAY)
    return p


def add_body_paragraph(doc, text, after=Pt(4.5)):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=after)
    r = p.add_run(text)
    set_run(r, size=BODY_SIZE, color=DARKGRAY)
    return p


def add_summary_segments(doc, segments, after=Pt(4.5)):
    """Bold-capable paragraph for Career Summary — segments: list of (text, bold) tuples.
    Always use this (never add_body_paragraph) for Career Summary so keyword bolding
    is never silently lost."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=after)
    for text, bold in segments:
        r = p.add_run(text)
        set_run(r, size=BODY_SIZE, color=DARKGRAY, bold=bold)
    return p


# ---------- Document build ----------
def build_resume():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Remove default spacing baked into the Normal style so nothing leaks through
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = LINE_SPACING

    # ---- Header ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=Pt(0), after=Pt(2))
    set_run(p.add_run("GAGAN DEEP GARG"), size=NAME_SIZE, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=Pt(0), after=Pt(5))
    set_run(
        p.add_run(
            "DIGITAL & INTEGRATED MARKETING LEADER (20+ YRS) | "
            "SCALING GROWTH & REVENUE"
        ),
        size=TITLE_SIZE, color=MIDGRAY, bold=True,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=Pt(0), after=Pt(3))
    set_run(
        p.add_run("Open to Remote/Relocation  •  +91 9871121860  •  gdgarg@gmail.com  •  linkedin.com/in/gdgarg"),
        size=SMALL_SIZE, color=DARKGRAY,
    )
    add_bottom_border(p, sz=8)

    # ---- Career Summary ----
    # (Was "Executive Highlights" — WRITING_GUIDE.md's canonical heading is
    # "Career Summary". Also restored keyword bolding via add_summary_segments,
    # which this section previously lacked.)
    add_section_heading(doc, "Career Summary")
    add_summary_segments(doc, [
        ("Digital & Integrated Marketing leader with ", False), ("20 years of experience", True),
        (" leading ", False), ("go-to-market strategy", True), (", ", False), ("growth marketing", True),
        (", ", False), ("demand generation", True), (", and ", False), ("digital marketing", True),
        (" across ", False), ("B2B SaaS", True), (" and technology organizations. Proven success "
         "partnering with Sales, Product, Customer Success, and executive leadership to accelerate "
         "pipeline growth, strengthen digital visibility, and deliver measurable revenue impact across "
         "the Americas, EMEA, and APJC.", False),
    ], after=Pt(4.5))
    add_summary_segments(doc, [
        ("Experienced building integrated digital marketing programs for enterprise decision-makers "
         "across B2B SaaS and technology environments. Combines strategic leadership with hands-on "
         "execution across ", False), ("digital marketing", True), (", ", False), ("paid media", True),
        (", ", False), ("demand generation", True), (", ", False), ("marketing analytics", True),
        (", automation, and ", False), ("AI-enabled marketing operations", True), (".", False),
    ], after=Pt(0))

    # ---- Core Skills ----
    add_section_heading(doc, "Core Skills")
    add_body_paragraph(
        doc,
        "Go-to-Market (GTM) Strategy  •  Account-Based Marketing (ABM)  •  Enterprise Demand "
        "Generation  •  Digital Marketing  •  Integrated Marketing  •  Brand Strategy & Positioning  •  "
        "Integrated Marketing & Communications  •  Thought Leadership  •  Executive Engagement  "
        "•  Regional & Field Marketing  •  Customer Marketing  •  Pipeline Acceleration  •  "
        "Sales & Marketing Alignment  •  Marketing Budget Ownership  •  Marketing Analytics & "
        "ROI  •  Marketing Technology Stack  •  AI-Enabled Marketing Operations  •  "
        "Cross-functional Leadership",
        after=Pt(0),
    )

    # ---- Career Highlights ----
    add_section_heading(doc, "Career Highlights")
    add_bullet(doc, [("Managed ", False), ("$4M+ global marketing investments", True), (" across enterprise demand generation, GTM, and integrated marketing programs supporting communications and technology portfolios.", False)])
    add_bullet(doc, [("Increased ", False), ("qualified pipeline creation by 25%", True), (" through account-based engagement, integrated demand generation, and regional GTM execution.", False)])
    add_bullet(doc, [("Improved ", False), ("lead-to-opportunity conversion by 20%", True), (" through Sales-Marketing alignment, lifecycle optimization, and customer engagement programs.", False)])
    add_bullet(doc, [("Increased ", False), ("qualified conversions by 22%", True), (" while reducing ", False), ("CPA by 17%", True), (" through audience segmentation, experimentation, and analytics-led optimization.", False)])
    add_bullet(doc, [("Delivered ", False), ("15x ROMI", True), (" across global enterprise marketing investments.", False)])
    add_bullet(doc, [("Improved ", False), ("marketing attribution visibility by 25%", True), (" through GA4 modernization, Salesforce integration, and executive reporting frameworks.", False)])
    add_bullet(doc, [("Led executive engagement initiatives across webinars, conferences, customer events, and regional field marketing programs supporting enterprise customer acquisition and expansion.", False)])

    # ---- Leadership Experience ----
    add_section_heading(doc, "Leadership Experience")

    add_role_header(doc, "GLOBAL ABM (AMERICAS, EUROPE & APJC) & MARKETING AUTOMATION", "CIVICA UK Ltd", "Aug 2026 – Present")
    add_bullet(doc, [("Drive marketing automation via Salesforce Marketing Cloud across Americas, Europe, and APJC, with focused scaling across North America and APAC.", False)])
    add_bullet(doc, [("Lead Demandbase adoption and ABM strategy for global account engagement.", False)])
    add_bullet(doc, [("Own programmatic and display advertising across the global digital portfolio.", False)])
    add_bullet(doc, [("Align paid media and pipeline strategy with Sales sprint motions.", False)])
    add_bullet(doc, [("Manage a ", False), ("2-member team", True), (" and agency partners across global digital execution.", False)])

    add_role_header(doc, "GLOBAL DIGITAL MARKETING (AMERICAS, EUROPE & APJC)", "Keysight Technologies Ltd", "Jul 2024 – Mar 2026")
    add_bullet(doc, [("Lead global GTM and enterprise demand generation initiatives supporting technology businesses across APAC, EMEA, and the Americas.", False)])
    add_bullet(doc, [("Own ", False), ("$3M+ annual marketing investments", True), (", aligning budget allocation with pipeline generation, customer acquisition, and regional growth priorities.", False)])
    add_bullet(doc, [("Partner with Sales, Product Marketing, and regional leadership to design integrated account-based and demand generation programs targeting enterprise technology buyers.", False)])
    add_bullet(doc, [("Direct integrated marketing across digital, web, webinars, email, CRM, paid media, automation, and customer engagement channels.", False)])
    add_bullet(doc, [("Improved qualified conversions by ", False), ("22%", True), (" while reducing CPA by ", False), ("17%", True), (" through experimentation, audience segmentation, and analytics-driven optimization.", False)])
    add_bullet(doc, [("Build executive dashboards and pipeline reporting frameworks using Salesforce, GA4, Domo, and Power BI to support investment decisions and business planning.", False)])
    add_bullet(doc, [("Drive AI-enabled workflow automation and marketing operations improvements to increase execution speed, reporting quality, and operational scalability.", False)])

    add_role_header(doc, "REGIONAL DIGITAL MARKETING (APJC)", "Cisco Systems Ltd", "Feb 2011 – Apr 2024")
    add_bullet(doc, [("Led regional marketing strategy and enterprise demand generation programs across ", False), ("14+ APAC markets", True), (" supporting networking, collaboration, cloud, security, and communications technology portfolios.", False)])
    add_bullet(doc, [("Managed ", False), ("$4M+ annual marketing investments", True), (" across ABM, digital demand generation, search, social, programmatic, webinars, CRM, and integrated customer engagement programs.", False)])
    add_bullet(doc, [("Developed regional GTM plans aligned with Sales, Product Marketing, Customer Success, and executive leadership priorities across enterprise and service provider ecosystems.", False)])
    add_bullet(doc, [("Increased qualified pipeline growth by ", False), ("25%", True), (" through integrated demand generation, account-based engagement, and regional marketing execution.", False)])
    add_bullet(doc, [("Improved lead-to-opportunity conversion by ", False), ("20%", True), (" through lifecycle marketing, sales alignment, and customer engagement optimization.", False)])
    add_bullet(doc, [("Led executive engagement initiatives including industry conferences, webinars, executive roundtables, customer events, and regional field marketing programs.", False)])
    add_bullet(doc, [("Directed agency partners, marketing vendors, analytics teams, and cross-functional stakeholders across multiple markets to deliver consistent GTM execution.", False)])
    add_bullet(doc, [("Modernized marketing measurement and executive reporting through GA4 and Salesforce integration, improving attribution visibility by ", False), ("25%", True), (".", False)])

    # ---- Projects ----
    # (NEW — WRITING_GUIDE.md requires a Projects section between Leadership
    # Experience and Earlier Experience. Every metric below is reused from
    # Career Highlights above, not newly introduced.)
    add_section_heading(doc, "Projects")
    add_bullet(doc, [("Enterprise Demand Engine: ", True), ("Built integrated demand generation and ABM programs, increasing ", False), ("qualified pipeline creation by 25%", True), (" through account-based engagement and regional GTM execution.", False)])
    add_bullet(doc, [("Media Portfolio Optimization: ", True), ("Delivered ", False), ("15x ROMI", True), (" while improving ", False), ("qualified conversions by 22%", True), (" and reducing ", False), ("CPA by 17%", True), (" through audience segmentation and analytics-led optimization.", False)])
    add_bullet(doc, [("Marketing Measurement Framework: ", True), ("Modernized GA4 and Salesforce-integrated executive reporting, improving ", False), ("marketing attribution visibility by 25%", True), (".", False)])
    add_bullet(doc, [("Customer Engagement & Lifecycle Marketing: ", True), ("Improved ", False), ("lead-to-opportunity conversion by 20%", True), (" through Sales-Marketing alignment and lifecycle optimization.", False)])

    # ---- Earlier Experience ----
    # (Was previously merged under "Leadership Experience" — WRITING_GUIDE.md
    # requires these two roles under their own "Earlier Experience" heading,
    # positioned after Projects.)
    add_section_heading(doc, "Earlier Experience")

    add_role_header(doc, "INTEGRATED MARKETING", "Network18 Media Ltd", "2009 – 2011")
    add_bullet(doc, [("Managed integrated marketing campaigns across television, print, digital, sponsorships, and events.", False)])
    add_bullet(doc, [("Developed marketing solutions aligned with client business objectives, audience engagement, and brand positioning.", False)])
    add_bullet(doc, [("Coordinated agencies, creative partners, operations, and sales teams to ensure successful campaign delivery.", False)])
    add_bullet(doc, [("Presented campaign performance, business insights, and strategic recommendations to enterprise clients and senior stakeholders.", False)])

    add_role_header(doc, "MEDIA MARKETING & ADVERTISEMENTS", "Bennett Coleman & Co Ltd (Times of India)", "2003 – 2008")
    add_bullet(doc, [("Managed integrated advertising and media campaigns across print, radio, outdoor, digital, and event platforms.", False)])
    add_bullet(doc, [("Supported strategic planning, campaign development, customer engagement, and media optimization initiatives.", False)])
    add_bullet(doc, [("Coordinated campaign reporting, proposal development, stakeholder communication, and client presentations.", False)])

    # ---- Technology Proficiency ----
    # (Was "Technology & Marketing Stack Proficiency" — renamed to match
    # WRITING_GUIDE.md's canonical heading. The non-canonical "Thought
    # Leadership & Executive Engagement" section that used to sit here has
    # been removed: it isn't part of the required section order, and its
    # bullets carried no quantified metrics so they couldn't qualify as
    # Projects either per WRITING_GUIDE.md's Projects rule.)
    add_section_heading(doc, "Technology Proficiency")
    add_sub_label(doc, "CRM & Marketing Automation", "Salesforce CRM | HubSpot | Oracle Eloqua | Marketo Engage")
    add_sub_label(doc, "ABM & Demand Generation", "LinkedIn Campaign Manager | Google Ads | DV360 | Meta Ads | Intent & Audience Targeting")
    add_sub_label(doc, "Analytics & Business Intelligence", "Google Analytics 4 (GA4) | Power BI | Tableau | Domo | Looker Studio")
    add_sub_label(doc, "AI & Workflow Automation", "ChatGPT | Claude | n8n | AI-Assisted Marketing Workflows")

    # ---- Certifications ----
    add_section_heading(doc, "Certifications")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(1.5))
    set_run(p.add_run("Artificial Intelligence Applications for Business Leaders"), bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(6))
    set_run(p.add_run("Outskill | Apr 2026"), color=MIDGRAY, italic=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(1.5))
    set_run(p.add_run("Google Marketing Platform Certification"), bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(0))
    set_run(p.add_run("Google"), color=MIDGRAY, italic=True)

    # ---- Education ----
    add_section_heading(doc, "Education")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(1.5))
    set_run(p.add_run("Masters in Business Administration (MBA) – Marketing & Finance"), bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(0))
    set_run(p.add_run("Army College of Materials Management, Jabalpur, India"), color=MIDGRAY, italic=True)

    return doc


if __name__ == "__main__":
    doc = build_resume()
    doc.save("Gagan_Deep_Garg_Resume.docx")
    print("done")
